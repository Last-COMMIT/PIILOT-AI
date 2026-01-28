
import cv2
import fitz
import numpy as np
from typing import List, Dict
from docx import Document
from pathlib import Path
from app.utils.logger import logger
from app.utils.image_loader import load_image
from app.services.file.audio_masking import audio_pii_service

# Helper for image loading (simulating app.utils.image_loader)
def load_image(image_path: str):
    return cv2.imread(image_path)

class Masker:
    """파일 마스킹 처리 (문서, 이미지, 오디오, 비디오 통합)"""
    def __init__(self, mask_char='*'):
        self.mask_char = mask_char
        logger.info("Masker 초기화")

    def mask_text(self, text: str, entities: List[Dict]) -> str:
        """텍스트의 PII를 마스킹"""
        if not entities:
            return text

        entities_sorted = sorted(entities, key=lambda x: x['start'], reverse=True)

        masked_text = text
        for entity in entities_sorted:
            start = entity['start']
            end = entity['end']
            pii_text = entity['text']
            mask_str = self.mask_char * len(pii_text)
            masked_text = masked_text[:start] + mask_str + masked_text[end:]

        return masked_text

    def mask_image(self, image_path: str, faces: List[Dict]) -> bytes:
        """
        이미지 마스킹 (얼굴 모자이크) - 팀원 구현 부분 복원
        
        Args:
            image_path: 이미지 파일 경로 또는 base64 인코딩된 이미지
            faces: 얼굴 위치 리스트 [{"x": int, "y": int, "width": int, "height": int}, ...]
            
        Returns:
            마스킹된 이미지 (bytes)
        """
        logger.info(f"이미지 마스킹: {len(faces)}개 얼굴")
        
        # 이미지 로드 (공통 유틸리티 사용)
        img = load_image(image_path)
        if img is None:
            print(f"이미지 로드 실패: {image_path}")
            return b""
        
        # 얼굴 마스킹 처리
        for face in faces:
            x = face.get("x", 0)
            y = face.get("y", 0)
            width = face.get("width", 0)
            height = face.get("height", 0)
            
            # 좌표 유효성 검사
            if x < 0 or y < 0 or width <= 0 or height <= 0:
                continue
            
            # 이미지 범위 내인지 확인
            img_height, img_width = img.shape[:2]
            x1 = max(0, x)
            y1 = max(0, y)
            x2 = min(img_width, x + width)
            y2 = min(img_height, y + height)
            
            if x2 > x1 and y2 > y1:
                # 얼굴 영역 추출 및 blur 처리
                face_region = img[y1:y2, x1:x2]
                if face_region.size > 0:
                    # 강한 blur 적용
                    blurred_face = cv2.blur(face_region, (100, 100))
                    img[y1:y2, x1:x2] = blurred_face
        
        # 이미지를 bytes로 변환
        _, encoded_img = cv2.imencode('.jpg', img)
        return encoded_img.tobytes()
    
    def mask_audio(self, audio_data: str, detected_items: list) -> bytes:
        """
        음성 마스킹 (개인정보 부분 음소거 또는 변조)
        
        Args:
            audio_data: 음성 파일 경로 또는 base64 인코딩된 오디오 데이터
            detected_items: 탐지된 개인정보 리스트
            
        Returns:
            마스킹된 오디오 (bytes)
        """
        try:
            logger.info(f"음성 마스킹 시작: {len(detected_items)}개 항목")
            
            # 오디오 마스킹 서비스 사용
            masked_bytes = audio_pii_service.mask_audio(audio_data, detected_items)
            
            logger.info("음성 마스킹 완료")
            return masked_bytes
            
        except Exception as e:
            logger.error(f"음성 마스킹 중 오류 발생: {str(e)}")
            raise e
    
    def mask_video(self, video_path: str, faces: list, audio_items: list, 
                   save_path: str = None) -> bytes:
        """
        영상 마스킹 (얼굴 모자이크 + 오디오 마스킹)
        
        Args:
            video_path: 비디오 파일 경로 또는 base64 인코딩된 비디오
            faces: 얼굴 위치 리스트 (사용하지 않음 - 자동 감지 사용)
            audio_items: 오디오 개인정보 항목 리스트
            save_path: 마스킹된 비디오를 저장할 경로 (None이면 프로젝트 루트의 tests 폴더에 저장)
            
        Returns:
            마스킹된 비디오 (bytes)
        """
        import base64
        import tempfile
        import os
        import subprocess
        import shutil
        from app.services.file.face_detector import YOLOFaceDetector
        from app.services.file.blur_censor import BlurCensor
        from app.services.file.video_processor import VideoProcessorEnhanced
        
        logger.info(f"영상 마스킹 시작: {len(faces)}개 얼굴 정보, {len(audio_items)}개 오디오 항목")
        
        # 비디오 파일 경로 처리 (base64 또는 파일 경로)
        input_video_path = None
        is_base64 = False
        output_is_temp = False
        
        # save_path가 지정되지 않으면 프로젝트 루트의 tests 폴더에 저장
        if save_path is None:
            # 프로젝트 루트 찾기 (app/services/file/masker.py 기준 상위 3단계)
            current_file = Path(__file__).resolve()
            project_root = current_file.parent.parent.parent.parent  # app/services/file -> app/services -> app -> root
            tests_dir = project_root / "tests"
            tests_dir.mkdir(exist_ok=True)
            
            # 타임스탬프 기반 파일명 생성
            import datetime
            timestamp = datetime.datetime.now().strftime('%Y%m%d_%H%M%S')
            save_path = str(tests_dir / f"masked_video_{timestamp}.mp4")
            output_is_temp = False  # 영구 저장
            logger.info(f"프로젝트 루트의 tests 폴더에 저장: {save_path}")
        
        # base64인지 확인
        if video_path.startswith("data:video") or len(video_path) > 1000:
            is_base64 = True
            # base64 디코딩하여 임시 파일로 저장
            try:
                if "," in video_path:
                    base64_data = video_path.split(",")[1]
                else:
                    base64_data = video_path
                
                video_bytes = base64.b64decode(base64_data)
                temp_input = tempfile.NamedTemporaryFile(delete=False, suffix='.mp4')
                temp_input.write(video_bytes)
                temp_input.close()
                input_video_path = temp_input.name
                logger.info("Base64 비디오를 임시 파일로 저장했습니다.")
            except Exception as e:
                logger.error(f"Base64 비디오 디코딩 오류: {e}")
                raise ValueError(f"비디오 디코딩 실패: {e}")
        else:
            # 파일 경로
            if not os.path.exists(video_path):
                raise FileNotFoundError(f"비디오 파일을 찾을 수 없습니다: {video_path}")
            input_video_path = video_path
        
        try:
            # 출력 파일 경로 설정 (save_path는 이미 설정됨)
            output_video_path = save_path
            # 디렉토리가 없으면 생성 (tests 폴더는 이미 생성되었지만 안전을 위해)
            output_dir = os.path.dirname(os.path.abspath(save_path))
            if output_dir:
                os.makedirs(output_dir, exist_ok=True)
            logger.info(f"마스킹된 비디오를 저장합니다: {save_path}")
            
            # 얼굴 감지기 초기화
            detector = YOLOFaceDetector(
                conf_threshold=0.25,
                iou_threshold=0.45,
                imgsz=640,
                enhance_image=True
            )
            
            # 블러 처리기 초기화
            blur_censor = BlurCensor(blur_factor=99)
            
            # 비디오 프로세서 초기화 (Kalman Filter 사용)
            processor = VideoProcessorEnhanced(
                detector=detector,
                censor=blur_censor,
                max_track_frames=10,
                iou_threshold=0.3,
                use_kalman=True
            )
            
            # 비디오 처리 (얼굴 블러 적용)
            logger.info("비디오 얼굴 마스킹 처리 중...")
            processor.process_video(input_video_path, output_video_path, conf_thresh=0.25)
            
            # 오디오 처리 (추출 → 마스킹 → 합성)
            final_video_path = output_video_path
            temp_audio_path = None
            temp_masked_audio_path = None
            temp_final_video_path = None
            
            try:
                # 비디오에서 오디오 추출
                temp_audio = tempfile.NamedTemporaryFile(delete=False, suffix='.mp3')
                temp_audio.close()
                audio_extract_path = temp_audio.name
                temp_audio_path = audio_extract_path
                
                # ffmpeg로 오디오 추출
                extract_cmd = [
                    'ffmpeg', '-i', input_video_path,
                    '-vn', '-acodec', 'libmp3lame',
                    '-y', audio_extract_path
                ]
                try:
                    subprocess.run(extract_cmd, check=True, capture_output=True)
                    logger.info("비디오에서 오디오 추출 완료")
                    
                    # 오디오 마스킹 처리
                    if audio_items:
                        logger.info(f"오디오 마스킹 처리 중... ({len(audio_items)}개 항목)")
                        masked_audio_bytes = audio_pii_service.mask_audio(audio_extract_path, audio_items)
                        
                        # 마스킹된 오디오를 임시 파일로 저장
                        temp_masked_audio = tempfile.NamedTemporaryFile(delete=False, suffix='.mp3')
                        temp_masked_audio.write(masked_audio_bytes)
                        temp_masked_audio.close()
                        masked_audio_path = temp_masked_audio.name
                        temp_masked_audio_path = masked_audio_path
                    else:
                        # 오디오 항목이 없으면 원본 오디오 사용
                        logger.info("오디오 마스킹 항목이 없어 원본 오디오를 사용합니다.")
                        masked_audio_path = audio_extract_path
                    
                    # 마스킹된 오디오를 비디오에 합성
                    temp_final_video = tempfile.NamedTemporaryFile(delete=False, suffix='.mp4')
                    temp_final_video.close()
                    final_video_with_audio = temp_final_video.name
                    temp_final_video_path = final_video_with_audio
                    
                    logger.info(f"오디오 합성 중: 비디오={output_video_path}, 오디오={masked_audio_path}")
                    
                    merge_cmd = [
                        'ffmpeg', '-i', output_video_path,
                        '-i', masked_audio_path,
                        '-c:v', 'copy', '-c:a', 'aac',
                        '-map', '0:v:0', '-map', '1:a:0',
                        '-shortest', '-y', final_video_with_audio
                    ]
                    subprocess.run(merge_cmd, check=True, capture_output=True)
                    logger.info("오디오 합성 완료")
                    
                    # 최종 파일을 원래 경로로 복사 (오디오가 합성된 경우)
                    if final_video_with_audio != output_video_path:
                        shutil.copy2(final_video_with_audio, output_video_path)
                        final_video_path = output_video_path
                    else:
                        final_video_path = final_video_with_audio
                    
                except subprocess.CalledProcessError as e:
                    logger.warning(f"오디오 처리 중 오류 발생 (비디오만 저장): {e}")
                    if e.stderr:
                        try:
                            logger.warning(f"ffmpeg 오류 출력: {e.stderr.decode('utf-8', errors='ignore')}")
                        except:
                            pass
                    # 오디오 처리 실패 시 비디오만 저장
                    final_video_path = output_video_path
                except FileNotFoundError:
                    logger.warning("ffmpeg가 설치되어 있지 않습니다. 오디오 없이 비디오만 저장됩니다.")
                    final_video_path = output_video_path
                except Exception as e:
                    logger.warning(f"오디오 처리 중 예상치 못한 오류 발생 (비디오만 저장): {e}")
                    final_video_path = output_video_path
            except Exception as e:
                logger.warning(f"오디오 처리 초기화 실패 (비디오만 저장): {e}")
                final_video_path = output_video_path
            finally:
                # 임시 파일 정리 (오디오 관련)
                if temp_audio_path and os.path.exists(temp_audio_path):
                    try:
                        os.unlink(temp_audio_path)
                    except:
                        pass
                if temp_masked_audio_path and os.path.exists(temp_masked_audio_path):
                    try:
                        os.unlink(temp_masked_audio_path)
                    except:
                        pass
                if temp_final_video_path and temp_final_video_path != final_video_path and os.path.exists(temp_final_video_path):
                    try:
                        os.unlink(temp_final_video_path)
                    except:
                        pass
            
            # 결과 비디오를 bytes로 읽기
            with open(final_video_path, 'rb') as f:
                masked_video_bytes = f.read()
            
            logger.info(f"영상 마스킹 완료: {len(masked_video_bytes)} bytes")
            
            # tests 폴더에 저장된 경우 저장 경로 로그 출력
            if not output_is_temp:
                logger.info(f"마스킹된 비디오가 저장되었습니다: {save_path}")
            
            return masked_video_bytes
            
        finally:
            # 임시 파일 정리
            # 입력 임시 파일 삭제 (base64인 경우)
            if is_base64 and input_video_path and os.path.exists(input_video_path):
                try:
                    os.unlink(input_video_path)
                    logger.debug(f"입력 임시 파일 삭제: {input_video_path}")
                except Exception as e:
                    logger.warning(f"입력 임시 파일 삭제 실패: {e}")
            
            # 출력 파일은 tests 폴더에 영구 저장되므로 삭제하지 않음

    def mask_pdf(self, pdf_path: str, text_blocks: List[Dict],
                 entities_per_block: List[List[Dict]],
                 output_path: str):
        """PDF의 PII 영역을 정밀하게 마스킹"""
        doc = fitz.open(pdf_path)

        for block, entities in zip(text_blocks, entities_per_block):
            if not entities:
                continue

            page_num = block['page']
            page = doc[page_num]
            
            # block['bbox']는 2배 해상도 기준이므로, 원본 PDF 좌표로 변환 (나누기 2)
            scale_factor = 2.0
            x0, y0, x1, y1 = block['bbox'][0][0], block['bbox'][0][1], block['bbox'][2][0], block['bbox'][2][1]
            
            clip_rect = fitz.Rect(x0 / scale_factor, y0 / scale_factor, 
                                  x1 / scale_factor, y1 / scale_factor)
            
            # PII 엔티티별로 마스킹
            for entity in entities:
                target_text = entity['text']
                hit_rects = page.search_for(target_text, clip=clip_rect)
                
                if hit_rects:
                    for rect in hit_rects:
                        page.draw_rect(rect, color=(0, 0, 0), fill=(0, 0, 0))
                else:
                    try:
                        pts = block['bbox']
                        bx0 = pts[0][0] / scale_factor
                        by0 = pts[0][1] / scale_factor
                        bx1 = pts[2][0] / scale_factor
                        by1 = pts[2][1] / scale_factor
                        
                        fallback_rect = fitz.Rect(bx0, by0, bx1, by1)
                        page.draw_rect(fallback_rect, color=(0, 0, 0), fill=(0, 0, 0))
                    except Exception as e:
                        print(f"  [Warning] Fallback masking failed: {e}")

        doc.save(output_path)
        doc.close()
        print(f"PDF 저장: {output_path}")

    def mask_docx(self, docx_path: str, paragraphs: List[Dict],
                  entities_per_para: List[List[Dict]],
                  output_path: str):
        """DOCX의 PII를 마스킹 (문단 및 표 지원)"""
        doc = Document(docx_path)

        for para_info, entities in zip(paragraphs, entities_per_para):
            if not entities:
                continue

            target_para = None
            try:
                p_type = para_info.get('type', 'para')
                if p_type == 'para':
                    if 'idx' in para_info:
                        target_para = doc.paragraphs[para_info['idx']]
                    elif 'paragraph_idx' in para_info:
                        target_para = doc.paragraphs[para_info['paragraph_idx']]      
                elif p_type == 'table':
                    table = doc.tables[para_info['table_idx']]
                    row = table.rows[para_info['row_idx']]
                    cell = row.cells[para_info['col_idx']]
                    target_para = cell.paragraphs[para_info['para_idx']]
                
                if target_para:
                    masked_text = self.mask_text(target_para.text, entities)
                    target_para.text = masked_text
                    
            except IndexError:
                continue
            except Exception as e:
                print(f"마스킹 중 오류 발생: {e}")
                continue

        doc.save(output_path)
        print(f"DOCX 저장: {output_path}")

    def mask_txt(self, txt_path: str, lines: List[Dict],
                 entities_per_line: List[List[Dict]],
                 output_path: str):
        """TXT의 PII를 마스킹"""
        with open(txt_path, 'r', encoding='utf-8') as f:
            all_lines = f.readlines()

        for line_info, entities in zip(lines, entities_per_line):
            if not entities:
                continue

            line_idx = line_info['line_idx']
            original_line = all_lines[line_idx]
            masked_text = self.mask_text(original_line.strip(), entities)
            all_lines[line_idx] = masked_text + '\n'

        with open(output_path, 'w', encoding='utf-8') as f:
            f.writelines(all_lines)

        print(f"TXT 저장: {output_path}")