import os
from typing import List, Dict, Tuple, Optional
import tempfile
import warnings
from app.core.logging import logger

warnings.filterwarnings('ignore')

class TextExtractor:
    """다양한 파일에서 텍스트 추출"""
    def __init__(self, use_gpu=True):
        # 지연 import: easyocr은 무거운 라이브러리
        import easyocr
        self.reader = easyocr.Reader(['ko', 'en'], gpu=use_gpu)

    def extract_from_image(self, image_path: str) -> List[Dict]:
        """이미지에서 텍스트 추출"""
        results = self.reader.readtext(image_path)
        results = self._combine_results(results)

        extracted = []
        for bbox, text, conf in results:
            extracted.append({
                'text': text,
                'bbox': bbox,
                'confidence': float(conf)
            })

        logger.debug(f"{len(extracted)}개 텍스트 블록 추출 (병합됨)")
        return extracted

    def _combine_results(self, raw_results, x_thresh_ratio=1.5, y_thresh_ratio=0.5):
        if not raw_results:
            return []
        def get_center_y(bbox):
            ys = [p[1] for p in bbox]
            return sum(ys) / 4
        def get_height(bbox):
            ys = [p[1] for p in bbox]
            return max(ys) - min(ys)
        sorted_res = sorted(raw_results, key=lambda r: (min(p[1] for p in r[0]), min(p[0] for p in r[0])))
        merged = []
        current = list(sorted_res[0])
        for next_item in sorted_res[1:]:
            next_item = list(next_item)
            curr_bbox, curr_text, curr_conf = current
            next_bbox, next_text, next_conf = next_item
            h1, h2 = get_height(curr_bbox), get_height(next_bbox)
            avg_h = (h1 + h2) / 2
            yc1, yc2 = get_center_y(curr_bbox), get_center_y(next_bbox)
            curr_x_end = max(p[0] for p in curr_bbox)
            next_x_start = min(p[0] for p in next_bbox)
            x_dist = next_x_start - curr_x_end
            same_line = abs(yc1 - yc2) < (avg_h * y_thresh_ratio)
            is_close = x_dist < (avg_h * x_thresh_ratio)
            if same_line and is_close:
                new_text = curr_text + " " + next_text
                xs = [p[0] for p in curr_bbox] + [p[0] for p in next_bbox]
                ys = [p[1] for p in curr_bbox] + [p[1] for p in next_bbox]
                new_bbox = [[min(xs), min(ys)], [max(xs), min(ys)], [max(xs), max(ys)], [min(xs), max(ys)]]
                current = [new_bbox, new_text, (curr_conf + next_conf) / 2]
            else:
                merged.append(current)
                current = next_item
        merged.append(current)
        return merged

    def extract_from_pdf(self, pdf_path: str) -> List[Dict]:
        """PDF에서 텍스트 추출"""
        try:
            import fitz  # PyMuPDF - lazy import
        except ImportError:
            logger.error("PyMuPDF (fitz) 모듈이 설치되지 않았습니다. 'pip install pymupdf'를 실행하세요.")
            raise ImportError("PyMuPDF (fitz) 모듈이 필요합니다.")
        
        logger.info(f"PDF 텍스트 추출: {pdf_path}")
        doc = fitz.open(pdf_path)
        all_extracted = []
        total_pages = len(doc)
        temp_dir = tempfile.gettempdir()

        for page_num in range(total_pages):
            page = doc[page_num]
            blocks = page.get_text("blocks")
            text_blocks = [b for b in blocks if b[6] == 0]
            total_text_len = sum(len(b[4].strip()) for b in text_blocks)
            
            if total_text_len > 50:
                page_dict = page.get_text("dict")
                blocks = page_dict.get("blocks", [])
                scale = 2.0
                
                for block in blocks:
                    if "lines" not in block:
                        continue
                    for line in block["lines"]:
                        line_text = "".join([span["text"] for span in line["spans"]])
                        if not line_text.strip():
                            continue
                        
                        lx0, ly0, lx1, ly1 = line["bbox"]
                        slx0, sly0, slx1, sly1 = lx0 * scale, ly0 * scale, lx1 * scale, ly1 * scale
                        line_bbox = [[slx0, sly0], [slx1, sly0], [slx1, sly1], [slx0, sly1]]
                        
                        all_extracted.append({
                            'page': page_num,
                            'text': line_text.strip(),
                            'bbox': line_bbox,
                            'confidence': 1.0
                        })
            else:
                logger.debug(f"페이지 {page_num+1}: 텍스트 레이어 부족 ({total_text_len}자) -> 이미지 변환 후 OCR 수행")
                pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))
                img_path = os.path.join(temp_dir, f"page_{page_num}.png")
                pix.save(img_path)

                results = self.reader.readtext(img_path)
                for bbox, text, conf in results:
                    all_extracted.append({
                        'page': page_num,
                        'text': text,
                        'bbox': bbox,
                        'confidence': conf
                    })

                if os.path.exists(img_path):
                    os.remove(img_path)

        doc.close()
        logger.info(f"✓ {len(all_extracted)}개 텍스트 블록 추출 (총 {total_pages}페이지)")
        return all_extracted

    def extract_from_docx(self, docx_path: str) -> List[Dict]:
        """DOCX에서 텍스트 추출 (문단 + 표)"""
        try:
            from docx import Document
        except ImportError:
            logger.error("python-docx 모듈이 설치되지 않았습니다. 'pip install python-docx'를 실행하세요.")
            raise ImportError("python-docx 모듈이 필요합니다.")
        
        logger.info(f"DOCX 텍스트 추출: {docx_path}")
        doc = Document(docx_path)
        extracted = []

        # 1. 본문 문단 추출
        for idx, para in enumerate(doc.paragraphs):
            if para.text.strip():
                extracted.append({
                    'type': 'para',
                    'idx': idx,
                    'text': para.text
                })

        # 2. 표(Table) 내 텍스트 추출
        for t_idx, table in enumerate(doc.tables):
            for r_idx, row in enumerate(table.rows):
                for c_idx, cell in enumerate(row.cells):
                    for p_idx, para in enumerate(cell.paragraphs):
                        if para.text.strip():
                            extracted.append({
                                'type': 'table',
                                'table_idx': t_idx,
                                'row_idx': r_idx,
                                'col_idx': c_idx,
                                'para_idx': p_idx,
                                'text': para.text
                            })

        logger.debug(f"{len(extracted)}개 텍스트 블록(문단+표) 추출")
        return extracted

    def extract_from_txt(self, txt_path: str) -> List[Dict]:
        """TXT에서 텍스트 추출"""
        logger.info(f"TXT 텍스트 추출: {txt_path}")
        with open(txt_path, 'r', encoding='utf-8') as f:
            lines = f.readlines()

        extracted = []
        for idx, line in enumerate(lines):
            if line.strip():
                extracted.append({
                    'text': line.strip(),
                    'line_idx': idx
                })

        logger.debug(f"{len(extracted)}개 라인 추출")
        return extracted