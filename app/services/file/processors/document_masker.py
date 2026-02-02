"""
문서 마스킹 처리 (PDF, DOCX, TXT, 텍스트)
(기존 masker.py에서 분리)
"""
from typing import List, Dict
from app.core.logging import logger
import fitz
import numpy as np
import cv2
from docx.opc.constants import RELATIONSHIP_TYPE as RT


class DocumentMasker:
    """문서 마스킹 처리"""

    def __init__(self, mask_char='*', face_detector=None, image_masker=None):
        self.mask_char = mask_char
        self.face_detector = face_detector
        self.image_masker = image_masker

    def mask_text(self, text: str, entities: List[Dict]) -> str:
        """텍스트의 PII를 마스킹"""
        if not entities:
            return text
        entities_sorted = sorted(entities, key=lambda x: x['start'], reverse=True)
        masked_text = text
        for entity in entities_sorted:
            start = entity['start']
            end = entity['end']
            pii_text = entity['text']
            mask_str = self.mask_char * len(pii_text)
            masked_text = masked_text[:start] + mask_str + masked_text[end:]
        return masked_text

    def mask_pdf(self, pdf_path: str, text_blocks: List[Dict],
                 entities_per_block: List[List[Dict]], output_path: str):
        """PDF의 PII 영역을 정밀하게 마스킹"""
        try:
            import fitz  # PyMuPDF - lazy import
        except ImportError:
            logger.error("PyMuPDF (fitz) 모듈이 설치되지 않았습니다. 'pip install pymupdf'를 실행하세요.")
            raise ImportError("PyMuPDF (fitz) 모듈이 필요합니다.")
        
        doc = fitz.open(pdf_path)
        for block, entities in zip(text_blocks, entities_per_block):
            if not entities:
                continue
            page_num = block['page']
            page = doc[page_num]
            scale_factor = 2.0
            x0, y0, x1, y1 = block['bbox'][0][0], block['bbox'][0][1], block['bbox'][2][0], block['bbox'][2][1]
            clip_rect = fitz.Rect(x0 / scale_factor, y0 / scale_factor,
                                  x1 / scale_factor, y1 / scale_factor)
            for entity in entities:
                target_text = entity['text']
                hit_rects = page.search_for(target_text, clip=clip_rect)
                if hit_rects:
                    for rect in hit_rects:
                        page.draw_rect(rect, color=(0, 0, 0), fill=(0, 0, 0))
                else:
                    try:
                        pts = block['bbox']
                        bx0 = pts[0][0] / scale_factor
                        by0 = pts[0][1] / scale_factor
                        bx1 = pts[2][0] / scale_factor
                        by1 = pts[2][1] / scale_factor
                        fallback_rect = fitz.Rect(bx0, by0, bx1, by1)
                        page.draw_rect(fallback_rect, color=(0, 0, 0), fill=(0, 0, 0))
                    except Exception as e:
                        logger.warning(f"Fallback masking failed: {e}")

        # 2. 이미지 마스킹 (얼굴 비식별화)
        if self.face_detector and self.image_masker:
            try:
                self._mask_pdf_images(doc)
            except Exception as e:
                logger.error(f"PDF 이미지 마스킹 중 오류 발생: {e}")

        doc.save(output_path)
        doc.close()
        logger.info(f"PDF 저장: {output_path}")

    def _mask_pdf_images(self, doc):
        """PDF 내 이미지 추출 및 얼굴 마스킹"""

        for page_num in range(len(doc)):
            page = doc[page_num]
            image_list = page.get_images()
            
            for img_info in image_list:
                xref = img_info[0]
                try:
                    # 이미지 추출
                    base_image = doc.extract_image(xref)
                    image_bytes = base_image["image"]
                    
                    # 얼굴 감지
                    nparr = np.frombuffer(image_bytes, np.uint8)
                    img = cv2.imdecode(nparr, cv2.IMREAD_COLOR)
                    
                    if img is None:
                        continue
                        
                    faces = self.face_detector.detect(img)
                    
                    if faces:
                        logger.info(f"PDF 이미지(xref={xref})에서 {len(faces)}개 얼굴 감지됨")
                        
                        # YOLO 결과 변환 (List -> List[Dict])
                        faces_dict = []
                        for face in faces:
                            # YOLO format: [x1, y1, x2, y2, conf]
                            if isinstance(face, list) or isinstance(face, np.ndarray):
                                x1, y1, x2, y2 = map(int, face[:4])
                                faces_dict.append({
                                    "x": x1,
                                    "y": y1,
                                    "width": x2 - x1,
                                    "height": y2 - y1
                                })
                        
                        # 마스킹 수행
                        masked_bytes = self.image_masker.mask_image(image_bytes, faces_dict)
                        
                        if masked_bytes:
                            # 이미지 교체
                            page.replace_image(xref, stream=masked_bytes)
                except Exception as e:
                    logger.warning(f"PDF 이미지(xref={xref}) 처리 중 오류: {e}")
                    continue

    def mask_docx(self, docx_path: str, paragraphs: List[Dict],
                  entities_per_para: List[List[Dict]], output_path: str):
        """DOCX의 PII를 마스킹 (문단 및 표 지원)"""
        try:
            from docx import Document
        except ImportError:
            logger.error("python-docx 모듈이 설치되지 않았습니다. 'pip install python-docx'를 실행하세요.")
            raise ImportError("python-docx 모듈이 필요합니다.")
        
        doc = Document(docx_path)
        for para_info, entities in zip(paragraphs, entities_per_para):
            if not entities:
                continue
            target_para = None
            try:
                p_type = para_info.get('type', 'para')
                if p_type == 'para':
                    if 'idx' in para_info:
                        target_para = doc.paragraphs[para_info['idx']]
                    elif 'paragraph_idx' in para_info:
                        target_para = doc.paragraphs[para_info['paragraph_idx']]
                elif p_type == 'table':
                    table = doc.tables[para_info['table_idx']]
                    row = table.rows[para_info['row_idx']]
                    cell = row.cells[para_info['col_idx']]
                    target_para = cell.paragraphs[para_info['para_idx']]
                if target_para:
                    masked_text = self.mask_text(target_para.text, entities)
                    target_para.text = masked_text
            except IndexError:
                continue
            except Exception as e:
                logger.warning(f"마스킹 중 오류 발생: {e}")
                continue

        # 2. 이미지 마스킹 (얼굴 비식별화)
        if self.face_detector and self.image_masker:
            try:
                self._mask_docx_images(doc)
            except Exception as e:
                logger.error(f"DOCX 이미지 마스킹 중 오류 발생: {e}")

        doc.save(output_path)
        logger.info(f"DOCX 저장: {output_path}")

    def _mask_docx_images(self, doc):
        """DOCX 내 이미지 추출 및 얼굴 마스킹"""
        
        # doc.part.rels를 순회하여 이미지 파트 찾기
        for rel in doc.part.rels.values():
            if "image" in rel.target_ref:
                try:
                    image_part = rel.target_part
                    image_bytes = image_part.blob
                    
                    # 얼굴 감지
                    nparr = np.frombuffer(image_bytes, np.uint8)
                    img = cv2.imdecode(nparr, cv2.IMREAD_COLOR)
                    
                    if img is None:
                        continue
                        
                    faces = self.face_detector.detect(img)
                    
                    if faces:
                        logger.info(f"DOCX 이미지에서 {len(faces)}개 얼굴 감지됨")
                        
                        # YOLO 결과 변환 (List -> List[Dict])
                        faces_dict = []
                        for face in faces:
                            # YOLO format: [x1, y1, x2, y2, conf]
                            if isinstance(face, list) or isinstance(face, np.ndarray):
                                x1, y1, x2, y2 = map(int, face[:4])
                                faces_dict.append({
                                    "x": x1,
                                    "y": y1,
                                    "width": x2 - x1,
                                    "height": y2 - y1
                                })
                                
                        # 마스킹 수행
                        masked_bytes = self.image_masker.mask_image(image_bytes, faces_dict)
                        
                        if masked_bytes:
                            # 이미지 교체 (private blob 직접 수정)
                            image_part._blob = masked_bytes
                except Exception as e:
                    logger.warning(f"DOCX 이미지 처리 중 오류: {e}")
                    continue

    def mask_txt(self, txt_path: str, lines: List[Dict],
                 entities_per_line: List[List[Dict]], output_path: str):
        """TXT의 PII를 마스킹"""
        with open(txt_path, 'r', encoding='utf-8') as f:
            all_lines = f.readlines()
        for line_info, entities in zip(lines, entities_per_line):
            if not entities:
                continue
            line_idx = line_info['line_idx']
            original_line = all_lines[line_idx]
            masked_text = self.mask_text(original_line.strip(), entities)
            all_lines[line_idx] = masked_text + '\n'
        with open(output_path, 'w', encoding='utf-8') as f:
            f.writelines(all_lines)
        logger.info(f"TXT 저장: {output_path}")
